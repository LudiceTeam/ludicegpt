import os
import zipfile
import tempfile
from docx import Document
import easyocr
from PIL import Image
import io
import shutil
import asyncio

async def extract_text_from_docx_with_images(docx_path: str) -> str:
  
    all_text = []
    
   
    try:
        doc = Document(docx_path)
        
       
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        
      
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    all_text.append(" | ".join(row_text))
    
    except Exception as e:
        all_text.append(f"[Ошибка чтения DOCX: {str(e)}]")
    
 
    images_text = await extract_text_from_docx_images(docx_path)
    if images_text:
        all_text.append("\n--- Текст с изображений ---")
        all_text.append(images_text)
    
    return "\n".join(all_text)


async def extract_text_from_docx_images(docx_path: str) -> str:
    temp_dir = tempfile.mkdtemp()
    images_text = []
    
    try:
       
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
          
            image_files = [f for f in docx_zip.namelist() 
                          if f.startswith('word/media/') and 
                          f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif'))]
            
            if not image_files:
                return ""
            
           
            reader = easyocr.Reader(['ru', 'en'])
            
            for i, image_path in enumerate(image_files):
               
                image_data = docx_zip.read(image_path)
                
               
                temp_image_path = os.path.join(temp_dir, f"image_{i}.png")
                with open(temp_image_path, 'wb') as f:
                    f.write(image_data)
                
              
                try:
                    result = reader.readtext(temp_image_path, detail=0, paragraph=True)
                    if result:
                        images_text.append(f"\n[Изображение {i+1}]:")
                        images_text.extend(result)
                except Exception as e:
                    images_text.append(f"\n[Ошибка OCR для изображения {i+1}: {str(e)}]")
    
    except Exception as e:
        images_text.append(f"[Ошибка извлечения изображений: {str(e)}]")
    
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
    
    return "\n".join(images_text)
